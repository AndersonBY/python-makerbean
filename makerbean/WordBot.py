# -*- coding: utf-8 -*-
# @Author: ander
# @Date:   2020-12-22 16:19:12
# @Last Modified by:   Anderson
# @Last Modified time: 2021-05-10 15:05:45
from docx import Document


class WordBot(object):
	"""docstring for WordBot"""

	def __init__(self):
		self.doc = Document()
		self.打开文档 = self.open
		self.保存文档 = self.save
		self.设置段落 = self.set_paragraph
		self.读取段落 = self.get_paragraph
		self.所有段落 = self.paragraphs
		self.清空内容 = self.clear

	def set_paragraph(self, index, text):
		if index < len(self.doc.paragraphs):
			self.doc.paragraphs[index].text = str(text)
			return True
		else:
			return False

	def get_paragraph(self, index):
		if index < len(self.doc.paragraphs):
			return str(self.doc.paragraphs[index].text)
		else:
			return ''

	@property
	def paragraphs(self):
		return [str(p.text) for p in self.doc.paragraphs]

	def clear(self):
		self.doc = Document()

	def open(self, filename):
		self.doc = Document(filename)

	def save(self, filename):
		if not filename:
			filename = 'tmp'
		self.doc.save(f'{filename}.docx')
